"""Генерация выгрузки в Word"""
import base64
import io
import os
import locale
from datetime import datetime
from decimal import Decimal
from typing import Optional
from docx.enum.section import WD_ORIENTATION
from docx.shared import Mm
from docx import Document

from ......api import models as m
from ......api.morpher import to_gent
from .....estimate import EstimateCalc, T_ADDITIONAL_COSTS_RTS
from ....queries.utils import transform_kwargs_by_mapped_spec
from ....queries.optim.sales.estimate.db_query import appendix_query, get_estimate_obj_by_filter_args
from .wordgen.word_styles import (
    decorate_table,
    add_styles,
)
from .wordgen.section_utils import add_page_number
from .wordgen.paragraph_utils import set_document_margins
from .wordgen.table_utils import (
    create_table_from_args,
    column_fmt,
    set_column_numbers,
)
from .wordgen.wordgen_utils import WordGeneratorContext


locale.setlocale(locale.LC_ALL, os.getenv('LOCALE'))


ADMIN_ISP_NAME = 'АО «РТС Деко»'
ADMIN_ISP_POSITION = 'Генерального директора'
ADMIN_ISP_FAMILY = 'Кагарова А.Г.'

MAPPED_KWARGS_APPENDIX_DOCX = {
    'appendix_code': ['appendix__code__exact', False],
    'appendix_id': ['id', True],
}


CELL_COLOR = 'F9F9F9'
CELL_COLOR_DARK = 'F3F3F3'
BORDER_COLOR = '#AEAEAE'


def tr_date_to_str(date_value: Optional[datetime]):
    return date_value.strftime('%d.%m.%Y') if date_value is not None else '__.__.20__'


def tr_price(val):
    if not val:
        return '-'
    return locale.format_string('%.2f', round(val, 2), grouping=True) if val is not None else ''


def sorter_additional_costs_rts(x):
    return x.start_period, x.end_period, x.title, x.city_title, x.count


def sorter_reservations_rts(x):
    return x.date_from, x.date_to, x.city_title, x.address_title, x.format_title


def mutate_generate_appendix_docx(parent, info, **kwargs):
    transform_kwargs_by_mapped_spec(MAPPED_KWARGS_APPENDIX_DOCX, kwargs, kwargs)
    if not kwargs:
        return None

    appendix = appendix_query(**kwargs)
    if not appendix:
        return None
    else:
        appendix = appendix[0]

    wg = WordGeneratorContext('тг.')
    self_company_info = m.SelfCompanyInfo.objects.all()[0]
    document = Document()
    project_cities_s = self_company_info.attachment_city

    text_appendix = f'Приложение №{appendix.code}'
    text_contract = 'к Договору №{0} от {1} г.,'.format(
        appendix.contract_code, tr_date_to_str(appendix.contract_registration_date)
    )

    text_city = 'город {0}, {1} г.'.format(project_cities_s, tr_date_to_str(appendix.created_date))

    add_styles(document, appendix.code, 'PTC Deco', 'Приложение', 'RU')

    for section in document.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.orientation = WD_ORIENTATION.LANDSCAPE
        new_width, new_height = section.page_height, section.page_width
        section.page_width = new_width
        section.page_height = new_height

        p = section.footer.add_paragraph()
        p.style = 'footer_style'
        r = p.add_run()
        add_page_number(r)

    set_document_margins(document, top_margin=1.27, bottom_margin=1.27, left_margin=1.27, right_margin=1.27)

    document.core_properties.title = f'{text_appendix} {text_contract} {text_city}'

    header = document.sections[0].header
    p_header = header.paragraphs[0]
    p_header.style = 'title_header'
    p_header.add_run(text_appendix)
    header.add_paragraph(text_contract).style = 'title_header'
    header.add_paragraph(text_city).style = 'title_header_last'

    p = document.add_paragraph(style='lemma_normal_top')
    p.add_run(f'Исполнитель: «{self_company_info.company_name}»,', style='lemma_bold')
    p.add_run(
        f' в лице {self_company_info.signatory_position_two} {self_company_info.signatory_two}, действующего на основании {self_company_info.document_base_two}, с одной стороны и'
    )

    p = document.add_paragraph(style='lemma_normal_2')
    p.add_run(f'Заказчик: «{appendix.client_title or "____"}»,', style='lemma_bold')
    p.add_run(
        f""" в лице {to_gent(appendix.contract_signatory_position) or "____"} {appendix.contract_signatory_two or "____"
              }, действующего на основании {to_gent(str(appendix.contract_based_on_document).lower()) or "____"
              }, с другой стороны, заключили настоящее Приложение к договору о нижеследующем"""
    )

    document.add_paragraph(
        f'Исполнитель в рамках настоящего Приложения оказывает нижеследующие услуги (бренд «{appendix.project_brand_title or "____"}»):',
        style='lemma_normal_3',
    )

    e: EstimateCalc = EstimateCalc(project_obj=appendix.project, appendix_obj=appendix)
    e.full_calc()
    e.calculate_address_program()

    address_program = []
    for i, ap_item in enumerate(e.address_programm):
        address_program.append(
            {
                '№ п/п': f'{i + 1}.',
                'Город': ap_item.city_title,
                'Формат': ap_item.format_title,
                'Период': f'{tr_date_to_str(ap_item.date_from)} ‑ {tr_date_to_str(ap_item.date_to)}',
                'Кол-во': str(ap_item.count),
                'Аренда': tr_price(ap_item.rent),
                'Скидка': f'{round(ap_item.discount_client_percent)}%' if ap_item.discount_client_percent else '-',
                'Аренда со скидкой': tr_price(ap_item.rent - ap_item.discount_client_percent),
                'Печать': tr_price(ap_item.printing),
                'Налог со скидкой': tr_price(ap_item.nalog - ap_item.nalog_discount_value),
                'Доп. расходы': tr_price(ap_item.additional),
                'Монтаж': tr_price(ap_item.mounting),
                'Итого': tr_price(ap_item.itog_summary),
            }
        )

    additional_costs = []
    ac_item: Optional[T_ADDITIONAL_COSTS_RTS]
    for i, ac_item in enumerate(sorted(e.additional_costs_rts, key=sorter_additional_costs_rts)):
        discount_summa = ac_item.discount_value * ac_item.count
        value_after_discount = ac_item.summa_before_discount - discount_summa
        discount = (
            (Decimal(100.0) * discount_summa / ac_item.summa_before_discount if ac_item.summa_before_discount else 0)
            if discount_summa
            else 0
        )
        additional_costs.append(
            {
                '№ п/п': f'{i + 1}.',
                'Наименование услуги': ac_item.title,
                'Город': ac_item.city_title,
                'Период': f'{tr_date_to_str(ac_item.start_period)} ‑ {tr_date_to_str(ac_item.end_period)}',
                'Кол-во': str(ac_item.count),
                'Стоимость': tr_price(ac_item.summa_before_discount),
                'Скидка': f'{round(discount)}%' if discount else '-',
                'Итого со скидкой': tr_price(value_after_discount),
            }
        )

    reservations = []
    for i, r_item in enumerate(sorted(e.reservations_rts, key=sorter_reservations_rts)):
        reservations.append(
            {
                '№ п/п': f'{i + 1}.',
                'Город': r_item.city_title,
                'Адрес': r_item.address_title,
                'Формат': r_item.format_title,
                'Период': f'{tr_date_to_str(r_item.date_from)} ‑ {tr_date_to_str(r_item.date_to)}',
            }
        )

    if not address_program and not additional_costs and not reservations:
        address_program.append(
            {
                '№ п/п': '1.',
                'Город': '',
                'Формат': '',
                'Период': '__.__.20__ ‑ __.__.20__',
                'Кол-во': '0',
                'Аренда': '',
                'Скидка': '',
                'Аренда со скидкой': '',
                'Печать': '',
                'Налог со скидкой': '',
                'Доп. расходы': '',
                'Монтаж': '',
                'Итого': '',
            }
        )
        additional_costs.append(
            {
                '№ п/п': f'1.',
                'Наименование услуги': '',
                'Город': '',
                'Период': '__.__.20__ ‑ __.__.20__',
                'Кол-во': '',
                'Стоимость': '',
                'Скидка': '-',
                'Итого со скидкой': ''
            })

        reservations.append(
            {
                '№ п/п': '1.',
                'Город': '',
                'Адрес': '',
                'Формат': '',
                'Период': '__.__.20__ ‑ __.__.20__',
            }
        )

    if address_program:
        table = create_table_from_args(
            document, CELLS_TABLE_ADDRESS_PROGRAM, address_program, True, 'adress_program_row', False
        )
        decorate_table(table, BORDER_COLOR, CELL_COLOR, cell_color_head=CELL_COLOR_DARK, nolast=True)


    if additional_costs:
        document.add_paragraph('Дополнительные расходы', style='additional_costs_header')
        table = create_table_from_args(
            document, CELLS_TABLE_ADDITIONAL_COSTS, additional_costs, True, 'adress_program_row', False
        )
        decorate_table(table, BORDER_COLOR, CELL_COLOR, cell_color_head=CELL_COLOR_DARK, nolast=True)

    p = document.add_paragraph('ИТОГО: ', style='additional_costs_header')
    p.add_run(wg.tr_price(e.itog.summary_estimate_value), style='lemma_bold')

    if reservations:
        document.add_paragraph(style='additional_costs_header')
        document.add_paragraph('Забронированные стороны', style='reservations_spacer')
        table = create_table_from_args(document, CELLS_TABLE_RESERVATIONS, reservations, True, 'adress_program_row', False)
        decorate_table(table, BORDER_COLOR, CELL_COLOR, cell_color_head=CELL_COLOR_DARK, nolast=True)

    return save_document(document)


def gen_cells_reservations_table():
    hdr = 'adress_program_header'
    rows = 'adress_program_row'
    row_r = 'adress_program_row_price'
    row_l = 'adress_program_row_left'
    return set_column_numbers(
        [
            column_fmt(0.87, [('№ п/п', hdr)], '№ п/п', rows),
            column_fmt(3.32, [('Город', hdr)], 'Город', rows),
            column_fmt(11.1,   [('Адрес', hdr)], 'Адрес', row_l),
            column_fmt(7.75, [('Формат', hdr)], 'Формат', rows),
            column_fmt(5.5,  [('Период', hdr)], 'Период', rows),
        ]
    )


CELLS_TABLE_RESERVATIONS = gen_cells_reservations_table()


def gen_cells_table_additional_cost():
    hdr = 'adress_program_header'
    rows = 'adress_program_row'
    row_r = 'adress_program_row_price'
    row_l = 'adress_program_row_left'
    return set_column_numbers(
        [
            column_fmt(0.87, [('№ п/п', hdr)], '№ п/п', rows),
            column_fmt(12.53,    [('Наименование услуги', hdr)], 'Наименование услуги', row_l),
            column_fmt(5.32, [('Город', hdr)], 'Город', rows),
            column_fmt(2.37, [('Период', hdr)], 'Период', rows),
            column_fmt(2,    [('Кол-во', hdr)], 'Кол-во', rows),
            column_fmt(2.15, [('Стоимость', hdr)], 'Стоимость', row_r),
            column_fmt(2.15, [('Скидка', hdr)], 'Скидка', rows),
            column_fmt(2.15, [('Итого', hdr), ('со скидкой', hdr)], 'Итого со скидкой', row_r),
        ]
    )


CELLS_TABLE_ADDITIONAL_COSTS = gen_cells_table_additional_cost()


def gen_cells_table_address_program():
    hdr = 'adress_program_header'
    rows = 'adress_program_row'
    row_r = 'adress_program_row_price'
    rows_l = 'adress_program_row_left'
    return set_column_numbers(
        [
            column_fmt(0.87, [('№ п/п', hdr)], '№ п/п', rows),
            column_fmt(3.32, [('Город', hdr)], 'Город', rows),
            column_fmt(4.66, [('Формат', hdr)], 'Формат', rows),
            column_fmt(2.37, [('Период', hdr)], 'Период', rows),
            column_fmt(1.23, [('Кол-во', hdr)], 'Кол-во', rows),
            column_fmt(2.14, [('Аренда', hdr)], 'Аренда', row_r),
            column_fmt(1.54, [('Скидка', hdr)], 'Скидка', rows),
            column_fmt(2.14, [('Аренда', hdr), ('со скидкой', hdr)], 'Аренда со скидкой', row_r),
            column_fmt(1.95, [('Печать', hdr)], 'Печать', row_r),
            column_fmt(1.75, [('Монтаж', hdr)], 'Монтаж', row_r),
            column_fmt(2.15, [('Доп.', hdr), ('работы', hdr)], 'Доп. расходы', row_r),
            column_fmt(2.27, [('Налог', hdr)], 'Налог со скидкой', row_r),
            column_fmt(2.15, [('Итого', hdr)], 'Итого', row_r),
        ]
    )


CELLS_TABLE_ADDRESS_PROGRAM = gen_cells_table_address_program()


def save_document(document):
    file_stream = io.BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    # if 0:
    #     fname = os.path.join('tmp', 'attachment.docx')
    #     with open(fname, 'wb') as f:
    #         f.write(file_stream.read())
    #     file_stream.seek(0)
    #     if os.name == 'nt':
    #         convert_to_pdf_win(os.path.join(os.getcwd(), fname))

    return base64.b64encode(file_stream.getvalue()).decode('ascii')


