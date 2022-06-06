from docx.enum.text import WD_LINE_SPACING, WD_BREAK
from docx.shared import Cm, Pt


def set_document_margins(document, top_margin=2.5, bottom_margin=2.5, left_margin=3, right_margin=1.5):
    """ Установить отступы полей в документе """
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(top_margin)
        section.bottom_margin = Cm(bottom_margin)
        section.left_margin = Cm(left_margin)
        section.right_margin = Cm(right_margin)


def add_marked_str(s, document, need_comma, style1, style2):
    p = document.add_paragraph(style=style1)
    p.add_run('-\t', style=style2)
    p.add_run(s + ',' if need_comma else s)


def upper_first(item):
    return item[0].upper() + item[1:]


def set_default_pformat(p, before=None, after=None, left_indent=None, tabstop=None, align=None):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    if before is not None and after is not None:
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
    if tabstop:
        p.paragraph_format.tab_stops.add_tab_stop(Cm(tabstop))
    if left_indent:
        p.paragraph_format.first_line_indent = Cm(left_indent)
    if align:
        p.alignment = align


def set_par_style(text, p, align, need_bold, style, before=None, after=None):
    r = p.add_run(text)
    p.style = style
    set_default_pformat(p, before=before, after=after)
    if align:
        p.paragraph_format.alignment = align
    if need_bold:
        r.font.bold = True
    return p


def add_breaked_line(text, add_run, *args):
    arr = text.split('<BR>')
    r = add_run(arr[0], *args)
    for item in arr[1:]:
        r.add_break(WD_BREAK.LINE)
        r = add_run(item, *args)
    return r


def add_breaked_paragraph(document, text, style, breaker):
    p = document.add_paragraph('', style=style)
    if breaker:
        p.paragraph_format.page_break_before = True
    return add_breaked_line(text, p.add_run)


def add_paragraph(document, text, style):
    return document.add_paragraph(text, style=style)
