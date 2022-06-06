import re
from typing import List, Tuple, Callable, Optional, Dict, Any

from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

from .list_utils import add_run_with_styled_tire, get_level_indices, add_signs_dot, make_text_item_with_tire
from .paragraph_utils import upper_first
from .wordgen_utils import VTOP, tr_descr


ISNUM = re.compile(r'^([0-9]+).*')


T_ADD_ROWCELL_FUN = Callable[[Any, str, str, str, bool], None]
T_CELL_TRANSFORM = Callable[[Dict[str, str], str], str]

T_HEAD_STR_STYLE = Tuple[str, Optional[str]]  # Текст строки, тип стиля

T_CELL_PARAM_ROW = Tuple[
    int,  # номер столбца
    List[T_HEAD_STR_STYLE],  # строки заголовка
    Tuple[
        str,  # Индекс в столбце данных - 0
        Optional[str],  # Стиль основных строк ячейки - 1
        Optional[str],  # Стиль первой строки ячейки - 2
        Optional[str],  # Стиль строки итогов (последней строки таблицы) - 3
        T_ADD_ROWCELL_FUN,  # Функция, выполняющая добавление текста в ячейку с форматированием
        # и установкой необходимых стилей - 4
        WD_CELL_VERTICAL_ALIGNMENT,  # Тип вертикального выравнивания - 5
        Optional[T_CELL_TRANSFORM],  # Функция-преобразователь строки - 6
        Optional[float],  # Ширина столбца - 7
    ],
]

T_CELL_PARAM_TYPE = List[T_CELL_PARAM_ROW]


def set_col_widths(table, widths):
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width


def set_col_widths_spec(table, spec: T_CELL_PARAM_TYPE):
    for _ in table.rows:
        for _, width in enumerate([x[7] for _, _, x in spec]):
            if width is None:
                return

    for row in table.rows:
        for idx, width in enumerate([x[7] for _, _, x in spec]):
            row.cells[idx].width = Cm(width)


def add_rowcell(cell, text, style, style_firstline, valign, ishead=None):
    """ Добавить ячейку строки в таблицу """
    if text is None:
        text = ''

    text_arr = [x.strip() for x in str(text).split('\n')]
    if text_arr:
        cell.text = ''
        p = cell.paragraphs[0]
        p.add_run(tr_descr(text_arr[0]))
        p.style = style

        for text_item in text_arr[1:]:
            p = cell.add_paragraph('')
            p.add_run(text_item)
            p.style = style

    cell.vertical_alignment = valign


def add_rowcell_work(cell, text, style, style_firstline, valign, ishead=False):
    """ Добавить широкую ячейку в строку таблицы с подуровневой стилизацией контента
    @param cell:
    @param text:
    @param style:
    @param style_firstline:
    @param valign:
    @param ishead: Текст в ячейке имеет внутренний заголовок или нет
    """
    global ISNUM
    if text is None:
        text = ''

    pt2_cnt = 10
    text_arr = [x.strip() for x in str(text).split('\n')]
    if text_arr:
        cell.text = ''
        p = cell.paragraphs[0]
        p.style = style_firstline

        if ishead:
            add_run_with_styled_tire(p, tr_descr(text_arr[0]), style_tire=style_firstline + '_tire')
        else:
            p.add_run(tr_descr(text_arr[0]))

        arr_it = text_arr[1:]
        max_i = len(arr_it) - 1
        start_maker = '-\t'

        level3_idx = 0
        level2_idx = 0

        for i, text_item in enumerate(arr_it):
            if text_item and text_item != 'None':
                text_item = tr_descr(text_item)
                if text_item[:2] != start_maker:
                    isnum = ISNUM.match(text_item)

                    num, level2_idx, level3_idx, separator = get_level_indices(isnum, level2_idx, level3_idx)

                    if level3_idx:
                        style = 'Tabcell_pt1_3lev' if level2_idx < pt2_cnt else 'Tabcell_pt2_3lev'
                    else:
                        style = 'Tabcell_pt1' if level2_idx < pt2_cnt else 'Tabcell_pt2'

                    text_item = upper_first(add_signs_dot(text_item))

                    text_item = '{0}.{1}{2}'.format(level2_idx, separator, text_item)
                else:
                    style, text_item = make_text_item_with_tire(
                        start_maker,
                        text_item,
                        level2_idx,
                        i,
                        arr_it,
                        max_i,
                        'Tabcell_tire_pt1',
                        'Tabcell_tire_pt2',
                        'Tabcell_tire_pt1_last',
                        'Tabcell_tire_pt2_last',
                    )

                cell.add_paragraph(text_item, style=style)

    cell.vertical_alignment = valign


def add_cell_paragraph_and_run(cell, style, text, need_new_paragraph, valign=None):
    if need_new_paragraph is None:
        p = cell.paragraphs[0]
        p.style = style
    else:
        p = cell.add_paragraph('', style=style)
    r = p.add_run(text)

    if valign:
        cell.vertical_alignment = valign
    return p, r


def create_table(document, rows, cols):
    """ Создать таблицу """
    table = document.add_table(rows=rows, cols=cols)
    ts = document.styles['Table Grid']
    ts.font.name = 'Arial'
    ts.font.size = Pt(10)
    table.style = ts
    return table


def column_fmt(
    col_width: Optional[float],
    column_header: List[T_HEAD_STR_STYLE],
    data_index: str,
    style_cell_first_str: Optional[str] = None,
    vertical_align: WD_CELL_VERTICAL_ALIGNMENT = WD_CELL_VERTICAL_ALIGNMENT.CENTER,
    style_row_itogs: Optional[str] = None,
    style_cell_internal_str: Optional[str] = None,
    fn_cell_transformer: Optional[T_CELL_TRANSFORM] = None,
    fn_add_rowcell: T_ADD_ROWCELL_FUN = add_rowcell,
) -> T_CELL_PARAM_ROW:
    """
        Создать спецификатор форматирования столбца таблицы для добавления в список спецификаторов

    @param col_width: ширина столбца
    @param column_number:  номер столбца
    @param column_header:  строки заголовка
    @param data_index: Индекс в столбце данных
    @param style_cell_internal_str:  Стиль основных строк ячейки
    @param style_cell_first_str:  Стиль первой строки ячейки
    @param style_row_itogs:  Стиль строки итогов (последней строки таблицы)
    @param fn_add_rowcell: Функция, выполняющая добавление текста в ячейку с форматированием и установкой необходимых
                           стилей
    @param vertical_align: Тип вертикального выравнивания
    @param fn_cell_transformer: Функция-преобразователь строки
    """
    return (
        0,
        column_header,
        (
            data_index,  # 0
            style_cell_first_str,  # 1
            style_cell_internal_str,  # 2
            style_row_itogs if style_row_itogs is not None else style_cell_first_str,  # 3
            fn_add_rowcell,  # 4
            vertical_align,  # 5
            fn_cell_transformer,  # 6
            col_width,  # 7
        ),
    )


def set_column_numbers(fmt_cells: T_CELL_PARAM_TYPE) -> T_CELL_PARAM_TYPE:
    return [(i, *x[1:]) for i, x in enumerate(fmt_cells)]


def create_table_from_args(document, cells_param: T_CELL_PARAM_TYPE, src: List[Dict[str, Any]],
                           minus_is_center=False, center_style=None, need_itogs=True, need_ishead=False):
    """ Создать таблицу и заполнить ее по спецификации
    @param document:
    @param cells_param: Форматный спецификатор столбцов
    @param src: Источник данных
    @param need_ishead: Ячейки должны содержать текст с внутренним заголовком или нет
    @param minus_is_center заменять стиль минуса на другой, если текст ячейки содержит только минус
    @param center_style стиль центированного минуса
    @param need_itogs Нужна строка итогов или нет
    @return: объект-таблица в документе
    """

    table = create_table(document, 1, len(cells_param))
    hdr_cells = table.rows[0].cells

    num: int
    l: T_HEAD_STR_STYLE
    for num, l, _ in cells_param:
        p = None
        cell = hdr_cells[num]
        for column_name, style in l:
            p, _ = add_cell_paragraph_and_run(cell, style, column_name, p, VTOP)

    sz = len(src) - 1
    for i, work in enumerate(src):
        row_cells = table.add_row().cells
        for (
            num,
            _,
            (key, style_tabcell, style_firstline, style_last_line, fun_add_rowcell, valign, converter, _),
        ) in cells_param:
            value_text = converter(work, key) if converter else work[key]
            if minus_is_center and value_text == '-':
                style_last_line = center_style
                style_tabcell = center_style

            if i == sz and need_itogs:
                fun_add_rowcell(row_cells[num], value_text, style_last_line, style_last_line, valign)
            else:
                fun_add_rowcell(
                    row_cells[num],
                    value_text,
                    style_tabcell,
                    style_firstline,
                    valign,
                    ishead=need_ishead,
                )

    set_col_widths_spec(table, cells_param)

    return table


def indent_table(table, indent):
    # noinspection PyProtectedMember
    tbl_pr = table._element.xpath('w:tblPr')
    if tbl_pr:
        e = OxmlElement('w:tblInd')
        e.set(qn('w:w'), str(indent))
        e.set(qn('w:type'), 'dxa')
        tbl_pr[0].append(e)


def add_tabhead(document, table, cell, text, style='Tabhead'):
    assert style is not None
    table[cell].text = text
    p = table[cell].paragraphs[0]
    p.style = document.styles[style]
    return p
